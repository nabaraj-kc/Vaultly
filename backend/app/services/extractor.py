import io
import os
import logging
from pypdf import PdfReader
import docx

class ExtractorService:
    @staticmethod
    def extract_text_from_bytes(file_bytes: bytes, filename: str) -> str:
        """Safely parses files by type based on extensions and byte formats."""
        ext = os.path.splitext(filename)[1].lower()
        
        # Enforce static analysis check: Limit file processing size to 10MB
        if len(file_bytes) > 10 * 1024 * 1024:
            raise ValueError("File exceeds maximum allowable size of 10MB.")

        if ext == ".txt":
            return ExtractorService._extract_txt(file_bytes)
        elif ext == ".pdf":
            return ExtractorService._extract_pdf(file_bytes)
        elif ext in [".docx", ".doc"]:
            return ExtractorService._extract_docx(file_bytes)
        else:
            raise ValueError(f"Unsupported file format: {ext}. Only PDF, DOCX, and TXT are supported.")

    @staticmethod
    def _extract_txt(file_bytes: bytes) -> str:
        try:
            return file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            # Try parsing with latin-1 if utf-8 fails
            return file_bytes.decode("latin-1")

    @staticmethod
    def _extract_pdf(file_bytes: bytes) -> str:
        try:
            pdf_file = io.BytesIO(file_bytes)
            reader = PdfReader(pdf_file)
            text_accumulator = []
            
            # Limit page extraction to avoid CPU locks on excessively large files
            max_pages = min(len(reader.pages), 50)
            for i in range(max_pages):
                page = reader.pages[i]
                page_text = page.extract_text()
                if page_text:
                    text_accumulator.append(page_text)
            
            extracted = "\n".join(text_accumulator)
            if not extracted.strip():
                raise ValueError("PDF contains no extractable text. It may be scanned or empty.")
            return extracted
        except Exception as e:
            logging.error(f"Error parsing PDF fallback: {e}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")

    @staticmethod
    def _extract_docx(file_bytes: bytes) -> str:
        try:
            docx_file = io.BytesIO(file_bytes)
            doc = docx.Document(docx_file)
            text_accumulator = []
            
            for para in doc.paragraphs:
                if para.text:
                    text_accumulator.append(para.text)
                    
            # Parse simple tables inside docx
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells if cell.text]
                    if row_text:
                        text_accumulator.append(" | ".join(row_text))

            extracted = "\n".join(text_accumulator)
            if not extracted.strip():
                raise ValueError("DOCX contains no extractable text.")
            return extracted
        except Exception as e:
            logging.error(f"Error parsing DOCX fallback: {e}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
